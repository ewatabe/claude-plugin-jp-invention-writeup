#!/usr/bin/env python3
"""技術説明書DOCXビルダ（v2）。

idea-expanded.md、請求項.txt、figure-list.md、figure-descriptions.md を統合し、
サンプル `~/patent/sample/明細書.docx` のセクション構成に近い技術説明書を
Wordファイルとして生成する。

これは「事務所が明細書を書くための土台」となる **テクニカルライト** であって、
正式な明細書そのものではない。事務所側で文体・章立てを整える前提。

v2 改善点:
- markdown 記法 (`**bold**`, `- bullet`, `# heading`, `*italic*`) を適切に処理
- Yu Gothic UI を latin / ea / cs / hAnsi すべての typeface に適用（日本語が他フォントで描画される問題を回避）
- セクション本文の段落分割を整える

Usage:
    python3 build_spec_docx.py \\
        --idea ~/patent/<案件名>/work/idea-expanded.md \\
        --claims ~/patent/<案件名>/output/請求項.txt \\
        --figure-list ~/patent/<案件名>/work/figure-list.md \\
        --figure-descriptions ~/patent/<案件名>/work/figure-descriptions.md \\
        --output ~/patent/<案件名>/output/技術説明書.docx
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

JP_FONT = "Yu Gothic UI"


# ============================================================
# フォント設定 — Word の East Asian typeface 対応
# ============================================================
def set_run_font_full(run, font_name=JP_FONT):
    """run の latin / east-asian / h-ansi / complex-script を全て font_name に設定。
    python-docx の `run.font.name = ...` は latin のみ設定するため、日本語が
    テーマ既定フォントにフォールバックしてしまう問題を回避する。
    """
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    for attr in ("w:ascii", "w:eastAsia", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), font_name)


# ============================================================
# Markdown パース
# ============================================================
_INLINE_RE = re.compile(r"\*\*(.+?)\*\*|__(.+?)__|`([^`]+)`")


def split_inline_markdown(text: str):
    """インラインのmarkdown (`**bold**`, `__bold__`, `` `code` ``) を分割。
    戻り値: [(text, format_dict)] のリスト。format_dict は {'bold': True, 'mono': True} 等。
    """
    parts = []
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], {}))
        if m.group(1):  # **bold**
            parts.append((m.group(1), {"bold": True}))
        elif m.group(2):  # __bold__
            parts.append((m.group(2), {"bold": True}))
        elif m.group(3):  # `code`
            parts.append((m.group(3), {"mono": True}))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], {}))
    return parts


def parse_md_line(line: str):
    """1行の markdown を (prefix_text, content_parts) に分解。
    - `- xxx`        → ("・ ", parts(xxx))
    - `1. xxx`       → ("1. ", parts(xxx))
    - `## xxx`       → None（明示的見出しはスキップ）
    - `xxx`          → ("", parts(xxx))
    """
    stripped = line.lstrip()
    if not stripped:
        return None
    if stripped.startswith("#"):
        return None
    m = re.match(r"^([-*])\s+(.*)$", stripped)
    if m:
        return ("・ ", split_inline_markdown(m.group(2)))
    m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
    if m:
        return (f"{m.group(1)}. ", split_inline_markdown(m.group(2)))
    return ("", split_inline_markdown(stripped))


def add_markdown_block(doc, md_text: str, *, size=11):
    """markdown ブロックをパースして docx 段落として追加。"""
    if not md_text or not md_text.strip():
        return
    for line in md_text.split("\n"):
        parsed = parse_md_line(line)
        if parsed is None:
            continue
        prefix, parts = parsed
        if not parts and not prefix:
            continue
        p = doc.add_paragraph()
        if prefix:
            r = p.add_run(prefix)
            set_run_font_full(r)
            r.font.size = Pt(size)
        for (text, fmt) in parts:
            if not text:
                continue
            r = p.add_run(text)
            set_run_font_full(r)
            r.font.size = Pt(size)
            if fmt.get("bold"):
                r.font.bold = True
            if fmt.get("italic"):
                r.font.italic = True


# ============================================================
# 段落番号 【００XX】
# ============================================================
class ParaCounter:
    def __init__(self, start=1):
        self.n = start

    def next(self) -> str:
        s = f"{self.n:04d}"
        full = s.translate(str.maketrans("0123456789", "０１２３４５６７８９"))
        self.n += 1
        return f"【{full}】"


# ============================================================
# 見出し / プレーン段落
# ============================================================
def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font_full(r)
    r.font.bold = True
    r.font.size = Pt(14 if level == 1 else 12)


def add_paragraph_plain(doc, text, *, bold=False, size=11):
    """markdown処理なしのプレーン段落（段落番号や符号説明など）。"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font_full(r)
    r.font.size = Pt(size)
    r.font.bold = bold
    return p


# ============================================================
# idea-expanded.md パース
# ============================================================
def parse_idea_md(idea_md: str) -> dict:
    """idea-expanded.md から各セクション本文を辞書で抜き出す。
    キーはセクション見出し（## の後の "N. xxx" の xxx 部分）。"""
    sections: dict[str, list[str]] = {}
    current = None
    for line in idea_md.splitlines():
        m = re.match(r"^##\s+\d+\.\s+(.+)$", line)
        if m:
            current = m.group(1).strip()
            sections[current] = []
        elif current:
            sections[current].append(line)
    return {k: "\n".join(v).strip() for k, v in sections.items()}


def extract_field(section_text: str, label: str) -> str | None:
    """セクション本文から「- **label**: 値」「label: 値」形式を抽出。"""
    pattern = re.compile(
        r"^[-*]?\s*\*?\*?" + re.escape(label) + r"\*?\*?\s*[:：]\s*(.+)$", re.MULTILINE
    )
    m = pattern.search(section_text)
    if m:
        return m.group(1).strip().lstrip("{").rstrip("}").strip()
    return None


# ============================================================
# figure-descriptions / figure-list パース
# ============================================================
def parse_figure_descriptions(md: str) -> tuple[list[tuple[int, str]], str]:
    blocks: list[tuple[int, str]] = []
    refs_text = ""
    current_num = None
    current_text: list[str] = []
    in_refs = False
    refs_lines: list[str] = []

    for line in md.splitlines():
        m_fig = re.match(r"^##\s+図(\d+)の説明", line)
        m_refs = re.match(r"^##\s+符号の説明", line)
        if m_fig:
            if current_num is not None and current_text:
                blocks.append((current_num, "\n".join(current_text).strip()))
            current_num = int(m_fig.group(1))
            current_text = []
            in_refs = False
        elif m_refs:
            if current_num is not None and current_text:
                blocks.append((current_num, "\n".join(current_text).strip()))
                current_num = None
                current_text = []
            in_refs = True
        else:
            if in_refs:
                refs_lines.append(line)
            elif current_num is not None:
                current_text.append(line)

    if current_num is not None and current_text:
        blocks.append((current_num, "\n".join(current_text).strip()))
    refs_text = "\n".join(refs_lines).strip()
    return blocks, refs_text


def parse_figure_list(md: str) -> list[tuple[int, str]]:
    out = []
    for line in md.splitlines():
        m = re.match(r"^\|\s*図(\d+)\s*\|\s*([^|]+)\s*\|", line)
        if m:
            out.append((int(m.group(1)), m.group(2).strip()))
    return out


# ============================================================
# メイン
# ============================================================
def build(args):
    idea_md = Path(args.idea).expanduser().read_text(encoding="utf-8") if args.idea else ""
    claims_text = Path(args.claims).expanduser().read_text(encoding="utf-8") if args.claims else ""
    figlist_md = Path(args.figure_list).expanduser().read_text(encoding="utf-8") if args.figure_list else ""
    figdesc_md = Path(args.figure_descriptions).expanduser().read_text(encoding="utf-8") if args.figure_descriptions else ""

    idea = parse_idea_md(idea_md)
    fig_list = parse_figure_list(figlist_md)
    fig_blocks, refs_text = parse_figure_descriptions(figdesc_md)

    doc = Document()
    counter = ParaCounter()

    # ---- 【書類名】明細書 ----
    add_heading(doc, "【書類名】 明細書", level=1)

    # ---- 【発明の名称】 ----
    name = ""
    if "発明の名称（仮）" in idea:
        name = extract_field(idea["発明の名称（仮）"], "名称") or ""
        # markdown装飾を除去
        name = re.sub(r"\*\*(.+?)\*\*", r"\1", name)
        name = re.sub(r"__(.+?)__", r"\1", name)
    add_heading(doc, f"【発明の名称】 {name}", level=2)

    # ---- 【技術分野】 ----
    add_heading(doc, "【技術分野】", level=2)
    add_paragraph_plain(doc, counter.next())
    if name:
        add_paragraph_plain(doc, f"本開示は、{name}の技術に関する。")
    else:
        add_paragraph_plain(doc, "本開示は、〇〇の技術に関する。")

    # ---- 【背景技術】 ----
    add_heading(doc, "【背景技術】", level=2)
    add_paragraph_plain(doc, counter.next())
    add_markdown_block(doc, idea.get("背景・現状の課題", "（背景を記載）"))

    # ---- 【先行技術文献】 ----
    add_heading(doc, "【先行技術文献】", level=2)
    add_heading(doc, "【特許文献】", level=3)
    add_paragraph_plain(doc, counter.next())
    add_paragraph_plain(doc, "【特許文献１】（公報番号を記載）")

    # ---- 【発明の概要】 ----
    add_heading(doc, "【発明の概要】", level=2)

    # 解決しようとする課題
    add_heading(doc, "【発明が解決しようとする課題】", level=3)
    add_paragraph_plain(doc, counter.next())
    add_markdown_block(doc, idea.get("本発明で解決したいこと", "（課題を記載）"))

    # 解決するための手段
    add_heading(doc, "【課題を解決するための手段】", level=3)
    add_paragraph_plain(doc, counter.next())
    add_markdown_block(doc, idea.get("解決手段の核（独立項の骨格）", "（解決手段を記載）"))
    add_paragraph_plain(doc, "その他の解決手段は実施形態中において適宜記載する。")

    # 発明の効果
    add_heading(doc, "【発明の効果】", level=3)
    add_paragraph_plain(doc, counter.next())
    add_markdown_block(doc, idea.get("効果", "（効果を記載）"))

    # ---- 【図面の簡単な説明】 ----
    add_heading(doc, "【図面の簡単な説明】", level=2)
    add_paragraph_plain(doc, counter.next())
    for num, title in fig_list:
        # 「を示す図である」がすでにある場合は重複させない
        if any(suffix in title for suffix in ("を示す図", "を示すフローチャート", "を示す表")):
            add_paragraph_plain(doc, f"【図{num}】{title}である。")
        else:
            add_paragraph_plain(doc, f"【図{num}】{title}を示す図である。")

    # ---- 【発明を実施するための形態】 ----
    add_heading(doc, "【発明を実施するための形態】", level=2)
    for num, body in fig_blocks:
        add_paragraph_plain(doc, counter.next())
        add_markdown_block(doc, body)

    # ---- 【符号の説明】 ----
    add_heading(doc, "【符号の説明】", level=2)
    add_paragraph_plain(doc, counter.next())
    if refs_text:
        for line in refs_text.splitlines():
            if line.strip():
                add_paragraph_plain(doc, line)

    # ---- 【書類名】特許請求の範囲 ----
    doc.add_page_break()
    add_heading(doc, "【書類名】 特許請求の範囲", level=1)
    if claims_text:
        for line in claims_text.splitlines():
            add_paragraph_plain(doc, line)

    # 保存
    out = Path(args.output).expanduser()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"OK: {out}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--idea", required=True)
    ap.add_argument("--claims", required=True)
    ap.add_argument("--figure-list", required=True)
    ap.add_argument("--figure-descriptions", required=True)
    ap.add_argument("--output", required=True)
    args = ap.parse_args()
    build(args)


if __name__ == "__main__":
    main()
